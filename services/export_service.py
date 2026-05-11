from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

class ExportService:
    @staticmethod
    def export_to_txt(report, file_path):
        content = f"{report.title}\n"
        content += "=" * len(report.title) + "\n\n"
        content += report.content if report.content else ""

        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)
        return True

    @staticmethod
    def export_to_docx(report, file_path):
        doc = DocxDocument()

        title = doc.add_heading(report.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if report.content:
            lines = report.content.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue

                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.", "10.")):
                    parts = line.split(".", 1)
                    if len(parts) == 2:
                        heading = doc.add_heading(f"{parts[0]}. {parts[1]}", level=2)
                else:
                    p = doc.add_paragraph(line)

        doc.save(file_path)
        return True

    @staticmethod
    def get_export_filename(default_name, extension):
        if extension == "txt":
            if not default_name.endswith(".txt"):
                return f"{default_name}.txt"
        elif extension == "docx":
            if not default_name.endswith(".docx"):
                return f"{default_name}.docx"
        return default_name